#!/usr/bin/env python3
"""
Generate cover letter for Vetcove - SaaS Implementation Manager.

Following CoverLetterGenerator.md workflow requirements:
- 250-400 words
- 3-6 paragraphs
- Company research incorporated
- ATS-friendly formatting
"""

from docx import Document
from docx.shared import Pt, Inches
from datetime import date
import os


def generate_cover_letter():
    """Generate Vetcove cover letter."""

    # Create document
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Date
    date_para = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    date_para.alignment = 0  # Left align

    # Recipient
    doc.add_paragraph()
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph("Vetcove")
    doc.add_paragraph("Remote")

    # Salutation
    doc.add_paragraph()
    salutation = doc.add_paragraph("Dear Hiring Manager,")

    # Opening paragraph
    doc.add_paragraph()
    opening = doc.add_paragraph(
        "I am writing to apply for the SaaS Implementation Manager position at Vetcove. "
        "With 5+ years delivering end-to-end platform implementations and client onboarding automation, "
        "I bring the technical depth and customer success orientation needed to ensure veterinary practices "
        "successfully adopt Vetcove's marketplace platform."
    )

    # Body paragraph 1: Requirements alignment
    doc.add_paragraph()
    body1 = doc.add_paragraph(
        "In my current role at Company A, I built a complete SaaS onboarding system across GoHighLevel CRM and Stripe, "
        "automating client workflows from discovery through go-live. This reduced onboarding time from 3-5 hours to "
        "1-2 hours per client while maintaining high-quality implementation standards. Previously at Company B, "
        "I led a digital transformation migrating 750+ client records to a Notion CRM platform with automated "
        "lifecycle stages and Zapier integrations. My hands-on experience with CRM configuration, data migration, "
        "and workflow automation directly aligns with Vetcove's need for technical implementation expertise. "
        "My Computer Science background enables me to troubleshoot technical issues independently "
        "and communicate effectively with both customers and engineering teams."
    )

    # Body paragraph 2: Company knowledge & fit
    doc.add_paragraph()
    body2 = doc.add_paragraph(
        "I am drawn to Vetcove's mission to modernize veterinary purchasing through technology. The opportunity "
        "to work with SMB veterinary practices resonates with my experience serving small business clients at "
        "Company B and Company A, where I managed end-to-end customer relationships and understood the importance "
        "of seamless platform adoption. Vetcove's B2B marketplace model requires both technical platform knowledge "
        "and strong consultative skills—strengths I've developed through my consulting background at a Big 4 firm and "
        "Consulting Firm combined with hands-on SaaS implementation work."
    )

    # Closing paragraph
    doc.add_paragraph()
    closing = doc.add_paragraph(
        "I would welcome the opportunity to discuss how my SaaS implementation experience, technical troubleshooting "
        "skills, and customer success focus can support Vetcove's veterinary practice clients. Thank you for your "
        "consideration. I am available at your convenience and look forward to speaking with you."
    )

    # Signature
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Your Name")

    # Save
    output_dir = "~/Career/Generated Assets/Vetcove"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Your_Name_CoverLetter_Vetcove.docx")

    doc.save(output_path)

    # Word count
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)

    return output_path, word_count


if __name__ == "__main__":
    try:
        print("=" * 70)
        print("GENERATING VETCOVE COVER LETTER")
        print("=" * 70)

        output_path, word_count = generate_cover_letter()

        print(f"\n✓ Cover letter generated successfully")
        print(f"✓ Output: {output_path}")
        print(f"✓ Word count: {word_count} words")
        print(f"✓ Target: 250-400 words")

        if 250 <= word_count <= 400:
            print("✓ Word count within target range")
        else:
            print("⚠ Word count outside target range")

    except Exception as e:
        print(f"\n✗ Error generating cover letter: {e}")
        import traceback
        traceback.print_exc()

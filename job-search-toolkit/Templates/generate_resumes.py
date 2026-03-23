#!/usr/bin/env python3
"""
Resume Tailoring Script for JobSearchToolkit
Generates tailored resumes for Vetcove, Process Street, Motorola, and HockeyStack roles
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Job data from job application files
JOBS = {
    "Vetcove": {
        "title": "SaaS Implementation Manager",
        "company": "Vetcove",
        "fit_score": 8.0,
        "salary": "$65K-$90K",
        "location": "Remote",
        "keywords": [
            "SaaS implementation", "platform building", "implementation lifecycle",
            "client onboarding", "technical troubleshooting", "discovery", "configuration",
            "training", "go-live", "SMB clients", "CRM", "system setup"
        ],
        "summary": "SaaS Implementation Manager with 5+ years of experience in platform implementation, client onboarding automation, and technical troubleshooting, delivering seamless go-live experiences through configuration, training, and end-to-end implementation lifecycle management in SMB-focused SaaS environments.",
        "skills_categories": {
            "Implementation & Onboarding": ["SaaS Implementation", "Client Onboarding Automation", "Go-Live Support", "Discovery & Requirements", "Configuration Management", "Training & Documentation"],
            "CRM & Platforms": ["GoHighLevel CRM", "Notion", "Stripe", "Calendly", "Zapier", "WordPress", "Trainerize"],
            "Process & Delivery": ["Implementation Lifecycle", "Technical Troubleshooting", "Data Migration", "Testing & UAT", "SOPs", "Change Management"],
            "Technical Skills": ["SQL", "Python", "API Integrations", "Workflow Automation", "System Architecture"],
            "Tools": ["Jira", "Confluence", "Excel", "Visio", "QuickBooks", "PowerPoint"]
        }
    },
    "Process_Street": {
        "title": "Junior Solutions Engineer",
        "company": "Process Street",
        "fit_score": 7.8,
        "salary": "$70K-$95K (est.)",
        "location": "Remote (Global)",
        "keywords": [
            "solutions engineering", "technical discovery", "demo building", "pre-sales",
            "no-code integrations", "Zapier", "API", "client-facing", "workflow automation",
            "technical communication", "requirements gathering"
        ],
        "summary": "Solutions Engineer with 4+ years of experience in technical discovery, workflow automation, and client-facing demonstrations, delivering solution implementations through no-code integrations, API connectivity, and requirements gathering in SaaS and workflow automation platforms.",
        "skills_categories": {
            "Solutions Engineering": ["Technical Discovery", "Demo Building", "Pre-Sales Support", "Requirements Gathering", "Solution Architecture", "Client Presentations"],
            "Workflow & Automation": ["Workflow Automation", "No-Code Integrations", "Zapier", "GoHighLevel", "ManyChat", "Process Automation"],
            "Technical Skills": ["API Integrations", "SQL", "Python", "System Configuration", "Technical Documentation", "Testing & UAT"],
            "Business Analysis": ["Requirements Elicitation", "BRD/FRD", "Process Mapping", "User Stories", "Stakeholder Management"],
            "Tools": ["Jira", "Confluence", "Notion", "Excel", "PowerPoint", "Calendly", "Stripe"]
        }
    },
    "Motorola": {
        "title": "IT Business Systems Analyst",
        "company": "Motorola Solutions",
        "fit_score": 7.8,
        "salary": "$118K-$121K",
        "location": "Remote",
        "keywords": [
            "business systems analyst", "IT systems", "SDLC", "UAT", "requirements documentation",
            "stakeholder management", "system integration", "enterprise IT", "technical specifications",
            "project coordination", "business requirements", "process improvement"
        ],
        "summary": "IT Business Systems Analyst with 5+ years of experience in enterprise IT transformations, SDLC project delivery, and requirements documentation, translating stakeholder needs into system delivery through UAT coordination, process mapping, and cross-functional collaboration in Fortune 500 technology environments.",
        "skills_categories": {
            "Business Analysis": ["Requirements Documentation", "BRD/FRD", "User Stories", "Acceptance Criteria", "Stakeholder Management", "Requirements Elicitation"],
            "Process & Methodology": ["SDLC", "Agile", "UAT Coordination", "Process Mapping", "As-Is/To-Be Analysis", "Change Management"],
            "IT Systems": ["System Integration", "Data Migration", "Technical Specifications", "Enterprise IT", "System Configuration", "Testing & Validation"],
            "Technical Skills": ["SQL", "Python", "API Integrations", "Data Analysis", "System Architecture", "Documentation"],
            "Tools": ["Jira", "Rally", "Confluence", "Visio", "Excel", "PowerPoint", "SharePoint"]
        }
    },
    "HockeyStack": {
        "title": "Implementation Manager",
        "company": "HockeyStack",
        "fit_score": 7.8,
        "salary": "$100K-$150K",
        "location": "Remote",
        "keywords": [
            "analytics implementation", "dashboard configuration", "data pipeline", "tracking setup",
            "marketing analytics", "attribution", "technical implementation", "B2B SaaS",
            "customer onboarding", "data-driven insights", "platform implementation"
        ],
        "summary": "Implementation Manager with 5+ years of experience in analytics platform implementation, dashboard configuration, and data pipeline setup, delivering actionable insights through tracking implementation, data analysis, and B2B SaaS client onboarding in marketing and revenue analytics environments.",
        "skills_categories": {
            "Analytics Implementation": ["Analytics Platform Implementation", "Dashboard Configuration", "Tracking Setup", "Data Pipeline Design", "Attribution Modeling", "Data Validation"],
            "Implementation & Onboarding": ["B2B SaaS Implementation", "Client Onboarding", "Technical Training", "Go-Live Support", "Documentation", "Configuration Management"],
            "Data & Analysis": ["Data Analysis", "SQL", "Python", "Excel Modeling", "Data Migration", "Data Quality Assurance"],
            "Technical Skills": ["API Integrations", "Workflow Automation", "System Integration", "Technical Troubleshooting", "Platform Configuration"],
            "Tools": ["Notion", "GoHighLevel", "Stripe", "Zapier", "Jira", "Confluence", "PowerPoint", "Google Analytics"]
        }
    }
}

# Professional background from Master Context
EXPERIENCE = {
    "Company A": {
        "title": "Business Systems Analyst",
        "dates": "Sep 2024 – Dec 2025",
        "location": "Chicago, IL",
        "bullets": {
            "default": [
                "Built end-to-end client onboarding workflow across GoHighLevel CRM and Stripe, automating waiver delivery, health questionnaire routing, and milestone tracking, reducing onboarding time from 3-5 hours to 1-2 hours per client",
                "Implemented Notion-based system of record for client delivery tracking and program management, centralizing 12 active client files with automated Calendly integration via Zapier for appointment synchronization",
                "Designed and launched Company A Lookbook Generator AI application (8,765 lines TypeScript/React), integrating Google Gemini API for personalized outfit recommendations with comprehensive Playwright test suite",
                "Performed UAT-style testing and defect triage on GoHighLevel automations, proactively identifying and resolving pipeline stage triggers, tag assignments, and email sequence failures before client impact"
            ]
        }
    },
    "Company B": {
        "title": "Business Operations Analyst",
        "dates": "Jun 2022 – Sep 2024",
        "location": "Chicago, IL",
        "bullets": {
            "default": [
                "Led digital transformation initiative, migrating 750+ client records from Google Sheets to Notion CRM with automated lifecycle stages, standardized data fields, and Calendly-Zapier integration for scheduling automation",
                "Generated $57,585 in solo revenue (22% of $256K company total) across 36 unique clients, managing end-to-end sales cycle from consultations to fittings and alterations",
                "Engineered website performance improvements achieving 75.55% increase in engaged sessions (1,878→3,281), 64.4% faster load time (6s→2s), and ranking #3 on Google for target keywords through UI redesign and SEO optimization",
                "Built social media presence from 0 to 20,460 followers across Instagram (18,855), Facebook (1,176), TikTok (347), and LinkedIn (82), driving 14,990% growth in brand visibility to 292,145 impressions"
            ]
        }
    },
    "Enterprise Corp": {
        "title": "Business Technology Analyst",
        "dates": "Oct 2020 – Mar 2022",
        "location": "Chicago, IL",
        "bullets": {
            "default": [
                "Delivered business analysis for Cisco EA 3.0 strategic transformation initiative (2 levels from CEO), producing daily executive progress reports, C-suite presentations, and workstream status tracking that contributed to 2.5x software growth and 3x renewal rates",
                "Managed team of 3 junior Business Analysts on HSBC wealth management engagement, assigning tasks, reviewing deliverables, and ensuring quality of requirements documentation and stakeholder communications",
                "Translated business requirements into technical specifications for Marriott UI/UX implementation, collaborating with PM and engineering teams using Confluence as documentation hub for cross-functional alignment"
            ]
        }
    },
    "Consulting_Firm": {
        "title": "Senior Consultant",
        "dates": "Mar 2022 – Jun 2022",
        "location": "Chicago, IL",
        "bullets": {
            "default": [
                "Contributed to AT&T customer retention project during strategic pivot, supporting cross-functional alignment across Product, Engineering, and Marketing teams with financial modeling and retention analysis"
            ]
        }
    }
}

EDUCATION = "University Name, B.S. Your Major, YYYY–2020"

CONTACT = {
    "name": "Your Name",
    "location": "Chicago, IL",
    "phone": "(555) 123-4567",
    "email": "user@example.com",
    "linkedin": "linkedin.com/in/yourprofile"
}


def create_tailored_resume(job_key, template_path, output_path):
    """Generate a tailored resume for a specific job."""
    job = JOBS[job_key]

    # Create new document
    doc = Document()

    # Set margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # HEADER - Name (centered, bold, 17pt)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(CONTACT["name"])
    name_run.bold = True
    name_run.font.name = 'Arial'
    name_run.font.size = Pt(17)
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)

    # CONTACT INFO (centered, 10pt)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{CONTACT['location']} | {CONTACT['phone']} | {CONTACT['email']} | {CONTACT['linkedin']}"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.name = 'Arial'
    contact_run.font.size = Pt(10)
    contact_para.paragraph_format.space_before = Pt(6)
    contact_para.paragraph_format.space_after = Pt(12)

    # PROFESSIONAL TITLE (centered, bold, 12pt)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(job["title"])
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(12)
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after = Pt(12)

    # SUMMARY (justified, 9.5pt)
    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_run = summary_para.add_run(job["summary"])
    summary_run.font.name = 'Arial'
    summary_run.font.size = Pt(9.5)
    summary_para.paragraph_format.space_before = Pt(6)
    summary_para.paragraph_format.space_after = Pt(12)

    # SKILLS SECTION
    skills_header = doc.add_paragraph()
    skills_header_run = skills_header.add_run("SKILLS")
    skills_header_run.bold = True
    skills_header_run.font.name = 'Arial'
    skills_header_run.font.size = Pt(11)
    skills_header.paragraph_format.space_before = Pt(6)
    skills_header.paragraph_format.space_after = Pt(6)

    for category, skills in job["skills_categories"].items():
        skills_para = doc.add_paragraph()
        # Category name (bold)
        cat_run = skills_para.add_run(f"{category}: ")
        cat_run.bold = True
        cat_run.font.name = 'Arial'
        cat_run.font.size = Pt(9.5)
        # Skills list (regular)
        skills_run = skills_para.add_run(", ".join(skills))
        skills_run.font.name = 'Arial'
        skills_run.font.size = Pt(9.5)
        skills_para.paragraph_format.space_before = Pt(2)
        skills_para.paragraph_format.space_after = Pt(2)

    # PROFESSIONAL EXPERIENCE
    exp_header = doc.add_paragraph()
    exp_header_run = exp_header.add_run("PROFESSIONAL EXPERIENCE")
    exp_header_run.bold = True
    exp_header_run.font.name = 'Arial'
    exp_header_run.font.size = Pt(11)
    exp_header.paragraph_format.space_before = Pt(12)
    exp_header.paragraph_format.space_after = Pt(6)

    # Add each job
    for job_name in ["Company A", "Company B", "Enterprise Corp", "Consulting_Firm"]:
        exp = EXPERIENCE[job_name]

        # Company name (bold, 10pt)
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(job_name.replace("_", " "))
        company_run.bold = True
        company_run.font.name = 'Arial'
        company_run.font.size = Pt(10)
        company_para.paragraph_format.space_before = Pt(6)
        company_para.paragraph_format.space_after = Pt(0)

        # Job title and dates (9.5pt)
        title_para = doc.add_paragraph()
        title_text = f"{exp['title']} | {exp['dates']} | {exp['location']}"
        title_run = title_para.add_run(title_text)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(9.5)
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(4)

        # Bullets (adjust count based on job)
        bullets = exp["bullets"]["default"]
        num_bullets = len(bullets) if job_name != "Consulting_Firm" else 1

        for i in range(num_bullets):
            bullet_para = doc.add_paragraph(style='List Bullet')
            bullet_run = bullet_para.add_run(bullets[i])
            bullet_run.font.name = 'Arial'
            bullet_run.font.size = Pt(9.5)
            bullet_para.paragraph_format.space_before = Pt(2)
            bullet_para.paragraph_format.space_after = Pt(2)
            bullet_para.paragraph_format.left_indent = Inches(0.25)

    # EDUCATION
    edu_header = doc.add_paragraph()
    edu_header_run = edu_header.add_run("EDUCATION")
    edu_header_run.bold = True
    edu_header_run.font.name = 'Arial'
    edu_header_run.font.size = Pt(11)
    edu_header.paragraph_format.space_before = Pt(12)
    edu_header.paragraph_format.space_after = Pt(6)

    edu_para = doc.add_paragraph()
    edu_run = edu_para.add_run(EDUCATION)
    edu_run.font.name = 'Arial'
    edu_run.font.size = Pt(9.5)

    # Save document
    doc.save(output_path)
    print(f"✓ Generated: {output_path}")
    return output_path


def main():
    """Generate all 4 tailored resumes."""
    template_path = "~/Downloads/Your_Name_Resume.docx"

    results = []
    for job_key, job_data in JOBS.items():
        company_name = job_data["company"].replace(" ", "_")
        output_path = f"~/Downloads/Your_Name_Resume_{company_name}.docx"

        print(f"\nGenerating resume for {job_data['company']} ({job_data['title']})...")
        create_tailored_resume(job_key, template_path, output_path)
        results.append({
            "company": job_data["company"],
            "title": job_data["title"],
            "fit_score": job_data["fit_score"],
            "file": output_path
        })

    # Print summary
    print("\n" + "="*80)
    print("RESUME GENERATION SUMMARY")
    print("="*80)
    for r in results:
        print(f"\n{r['company']} - {r['title']}")
        print(f"  Fit Score: {r['fit_score']}/10")
        print(f"  File: {r['file']}")

    print("\n✓ All resumes generated successfully!")
    print("\nNext step: Generate cover letters using CoverLetterGenerator workflow")


if __name__ == "__main__":
    main()

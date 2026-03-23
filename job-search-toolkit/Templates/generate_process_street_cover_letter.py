#!/usr/bin/env python3
"""
Generate cover letter for Process Street - Junior Solutions Engineer.
"""

from docx import Document
from docx.shared import Pt, Inches
from datetime import date
import os


def generate_cover_letter():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Date & Header
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph("Process Street")
    doc.add_paragraph("Remote")
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")

    # Opening
    doc.add_paragraph()
    doc.add_paragraph(
        "I am writing to apply for the Junior Solutions Engineer position at Process Street. "
        "With 5+ years combining technical discovery, client-facing communication, and workflow automation "
        "expertise, I am excited to bring my consultative approach and no-code integration skills to help "
        "Process Street customers implement workflow automation solutions."
    )

    # Body 1
    doc.add_paragraph()
    doc.add_paragraph(
        "At Company A, I built end-to-end workflow automations using no-code platforms (GoHighLevel, Zapier) "
        "and API integrations, connecting CRM, payment processing, and scheduling systems to automate client "
        "onboarding workflows. I also developed a full-stack TypeScript/React application integrating the "
        "Google Gemini API, demonstrating my ability to understand technical requirements and translate them "
        "into working solutions. My Enterprise Corp consulting background taught me to conduct technical discovery "
        "sessions with stakeholders, gather requirements, and present solutions that align business needs with "
        "technical capabilities. At Company B, I served 36 clients directly, managing consultative sales "
        "conversations and delivering customized solutions—experience that translates directly to pre-sales "
        "demos and technical discovery calls."
    )

    # Body 2
    doc.add_paragraph()
    doc.add_paragraph(
        "I am drawn to Process Street's mission to help teams standardize and automate their workflows. "
        "As someone who has built workflow automation systems from scratch and understands the challenges "
        "of process standardization, I am excited to help customers unlock the full value of the Process "
        "Street platform. The Junior Solutions Engineer role offers the perfect opportunity to combine my "
        "technical skills (API integrations, no-code tools) with my client-facing experience to support "
        "customers through their automation journey."
    )

    # Closing
    doc.add_paragraph()
    doc.add_paragraph(
        "I would welcome the opportunity to discuss how my technical discovery, workflow automation, and "
        "client communication skills can contribute to Process Street's customer success. Thank you for "
        "your consideration. I am available at your convenience and look forward to speaking with you."
    )

    # Signature
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Your Name")

    # Save
    output_dir = "~/Career/Generated Assets/Process Street"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Your_Name_CoverLetter_Process_Street.docx")
    doc.save(output_path)

    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    return output_path, word_count


if __name__ == "__main__":
    try:
        print("=" * 70)
        print("GENERATING PROCESS STREET COVER LETTER")
        print("=" * 70)

        output_path, word_count = generate_cover_letter()

        print(f"\n✓ Cover letter generated successfully")
        print(f"✓ Output: {output_path}")
        print(f"✓ Word count: {word_count} words")

        if 250 <= word_count <= 400:
            print("✓ Word count within target range")
        else:
            print("⚠ Word count outside target range")

    except Exception as e:
        print(f"\n✗ Error: {e}")
        import traceback
        traceback.print_exc()

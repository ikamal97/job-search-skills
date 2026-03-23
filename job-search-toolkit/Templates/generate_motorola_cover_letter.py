#!/usr/bin/env python3
"""
Generate cover letter for Motorola Solutions - IT Business Systems Analyst.
"""

from docx import Document
from docx.shared import Pt, Inches
from datetime import date
import os


def generate_cover_letter():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Date & Header
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph("Motorola Solutions")
    doc.add_paragraph("Remote")
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")

    # Opening
    doc.add_paragraph()
    doc.add_paragraph(
        "I am writing to apply for the IT Business Systems Analyst position at Motorola Solutions. "
        "With 5+ years delivering enterprise technology solutions through SDLC methodologies, requirements "
        "documentation, UAT coordination, and stakeholder management across Fortune 500 clients, I bring "
        "the technical foundation and business analysis expertise needed to support Motorola's IT transformation "
        "initiatives."
    )

    # Body 1
    doc.add_paragraph()
    doc.add_paragraph(
        "At Deloitte Consulting, I delivered business analysis for the Cisco EA 3.0 strategic IT transformation, "
        "a high-visibility initiative 2 levels from the CEO. I authored BRDs and FRDs, coordinated UAT activities, "
        "and produced daily executive status reports that supported 2.5x software growth and 3x partner renewal "
        "rates. I also managed a team of 3 junior Business Analysts on an HSBC wealth management enterprise system "
        "implementation, reviewing requirements documentation and ensuring SDLC compliance across Agile sprints. "
        "My Computer Science degree from Northwestern provides the technical foundation to understand complex "
        "system architectures and translate them into business requirements. Most recently at Kavalier, I authored "
        "technical specifications for CRM automation platforms, defined data models and workflow logic, and "
        "executed UAT-style testing with systematic defect tracking."
    )

    # Body 2
    doc.add_paragraph()
    doc.add_paragraph(
        "I am drawn to Motorola Solutions' role in public safety and enterprise communications technology. "
        "The IT Business Systems Analyst position offers the opportunity to apply my enterprise IT experience "
        "to mission-critical systems that protect communities and enable first responders. My background in "
        "large-scale IT transformations at Deloitte, combined with my technical Computer Science foundation, "
        "positions me well to support Motorola's complex system integrations and business requirements analysis."
    )

    # Closing
    doc.add_paragraph()
    doc.add_paragraph(
        "I would welcome the opportunity to discuss how my SDLC expertise, requirements documentation skills, "
        "and enterprise IT background can contribute to Motorola Solutions' technology initiatives. Thank you "
        "for your consideration. I am available at your convenience and look forward to speaking with you."
    )

    # Signature
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Idrees Kamal")

    # Save
    output_dir = "/Users/idreeskamal/Documents/Obsidian Vault/04 - Career/Generated Assets/Motorola Solutions"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Idrees_Kamal_CoverLetter_Motorola_Solutions.docx")
    doc.save(output_path)

    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    return output_path, word_count


if __name__ == "__main__":
    try:
        print("=" * 70)
        print("GENERATING MOTOROLA SOLUTIONS COVER LETTER")
        print("=" * 70)

        output_path, word_count = generate_cover_letter()

        print(f"\n✓ Cover letter generated successfully")
        print(f"✓ Output: {output_path}")
        print(f"✓ Word count: {word_count} words")

        if 250 <= word_count <= 400:
            print("✓ Word count within target range")
        else:
            print("⚠ Word count outside target range")

    except Exception as e:
        print(f"\n✗ Error: {e}")
        import traceback
        traceback.print_exc()

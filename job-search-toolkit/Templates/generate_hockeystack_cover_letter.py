#!/usr/bin/env python3
"""
Generate cover letter for HockeyStack - Implementation Manager.
"""

from docx import Document
from docx.shared import Pt, Inches
from datetime import date
import os


def generate_cover_letter():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Date & Header
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph("HockeyStack")
    doc.add_paragraph("Remote")
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")

    # Opening
    doc.add_paragraph()
    doc.add_paragraph(
        "I am writing to apply for the Implementation Manager position at HockeyStack. With 5+ years "
        "delivering analytics platform implementations and data-driven solutions, from tracking setup and "
        "dashboard configuration to go-live support, I bring the technical expertise and stakeholder management "
        "skills needed to ensure B2B customers successfully adopt HockeyStack's marketing analytics platform."
    )

    # Body 1
    doc.add_paragraph()
    doc.add_paragraph(
        "At Kavalier, I implemented an end-to-end analytics and tracking system across GoHighLevel CRM and "
        "Stripe, configuring conversion tracking, milestone metrics, and revenue reporting to monitor client "
        "lifecycle and business performance. I built data pipeline integrations connecting multiple platforms "
        "via API, establishing automated data flow with real-time synchronization. At Deloitte, I built executive "
        "dashboards and scenario models for Cisco's strategic transformation, tracking KPIs across multiple "
        "workstreams and presenting data-driven insights to C-suite stakeholders. My SQL knowledge and analytical "
        "background enable me to understand customer data requirements, troubleshoot tracking implementations, "
        "and configure dashboards that deliver actionable insights for marketing and revenue teams."
    )

    # Body 2
    doc.add_paragraph()
    doc.add_paragraph(
        "I am drawn to HockeyStack's mission to provide unified analytics for B2B marketing and revenue teams. "
        "As someone who has built analytics infrastructure for my own businesses and consulted on data-driven "
        "initiatives at Oliver Wyman and Deloitte, I understand the challenges marketing teams face with "
        "attribution and the value of connecting disparate data sources into a single source of truth. The "
        "Implementation Manager role offers the perfect blend of technical platform work and customer-facing "
        "consultation—both areas where I have demonstrated success."
    )

    # Closing
    doc.add_paragraph()
    doc.add_paragraph(
        "I would welcome the opportunity to discuss how my analytics implementation experience, technical "
        "troubleshooting skills, and data pipeline expertise can support HockeyStack's B2B customers. Thank "
        "you for your consideration. I am available at your convenience and look forward to speaking with you."
    )

    # Signature
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Idrees Kamal")

    # Save
    output_dir = "/Users/idreeskamal/Documents/Obsidian Vault/04 - Career/Generated Assets/HockeyStack"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Idrees_Kamal_CoverLetter_HockeyStack.docx")
    doc.save(output_path)

    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    return output_path, word_count


if __name__ == "__main__":
    try:
        print("=" * 70)
        print("GENERATING HOCKEYSTACK COVER LETTER")
        print("=" * 70)

        output_path, word_count = generate_cover_letter()

        print(f"\n✓ Cover letter generated successfully")
        print(f"✓ Output: {output_path}")
        print(f"✓ Word count: {word_count} words")

        if 250 <= word_count <= 400:
            print("✓ Word count within target range")
        else:
            print("⚠ Word count outside target range")

    except Exception as e:
        print(f"\n✗ Error: {e}")
        import traceback
        traceback.print_exc()
